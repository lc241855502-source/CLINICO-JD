from docx import Document
from docx.shared import Pt, Cm
import io
import os

def generate_jd_document(data, template_path):
    """
    基于公司真实模板填充JD内容
    通过关键词匹配定位填充位置，保证格式100%与原模板一致
    """
    if not template_path or not os.path.exists(template_path):
        raise FileNotFoundError(f"模板文件不存在: {template_path}")
    
    doc = Document(template_path)
    
    # 1. 填充表格中的基本信息字段
    _fill_table_fields(doc, data)
    
    # 2. 填充段落内容（岗位目的、职责、任职要求）
    _fill_paragraph_sections(doc, data)
    
    # 保存到内存
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def _fill_table_fields(doc, data):
    """填充表格中的字段
    短字段（基本信息）：填右边单元格
    长内容（岗位目的/职责/任职要求）：填标题下方的白色内容框
    """
    
    # 短字段映射：左标签，填右边值
    short_field_mapping = [
        (['Position Title', '职务名称', '职位名称', '岗位名称'], 'position_title'),
        (['Department', '部门', '所属部门'], 'department'),
        (['Line Manager', '直线经理', '汇报对象', 'Report to'], 'line_manager'),
        (['Subordinate', '下属', '下属人数'], 'subordinate'),
        (['Location', '工作地点', '地点'], 'location'),
    ]
    
    # 长内容字段映射：上标题，填下方内容框
    content_field_mapping = [
        (['Position purpose', '岗位目的', '职位概述', 'Position Summary'], 'position_purpose'),
        (['Major tasks', '主要职责', '岗位职责', 'Responsibilities', '工作描述'], 'responsibilities'),
        (['Qualifications', '任职要求', '岗位要求', 'Requirements'], 'qualifications'),
    ]
    
    for table in doc.tables:
        rows_count = len(table.rows)
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                
                # 1. 短字段：填右边格
                matched_short = None
                for keywords, field_name in short_field_mapping:
                    for kw in keywords:
                        if kw.lower() in cell_text.lower() and len(cell_text) < 30:
                            matched_short = field_name
                            break
                    if matched_short:
                        break
                
                if matched_short:
                    value = data.get(matched_short, '')
                    if value and col_idx + 1 < len(row.cells):
                        target_cell = row.cells[col_idx + 1]
                        _set_cell_content(target_cell, str(value))
                    continue
                
                # 2. 长内容：往下找真正的内容框（跳过合并的标题行）
                matched_content = None
                for keywords, field_name in content_field_mapping:
                    for kw in keywords:
                        if kw.lower() in cell_text.lower():
                            matched_content = field_name
                            break
                    if matched_content:
                        break
                
                if matched_content:
                    # 往下找内容框：跳过所有标题合并行，找到第一个空内容单元格
                    content_row = None
                    for r in range(row_idx + 1, min(row_idx + 8, rows_count)):
                        next_cell = table.rows[r].cells[col_idx]
                        next_text = next_cell.text.strip()
                        # 如果是空单元格，就是内容框了
                        if not next_text:
                            content_row = r
                            break
                        # 如果还包含标题关键词，说明还是合并的标题行，继续跳过
                        still_title = False
                        for kw in ['purpose', '岗位目的', 'task', '主要职责', '职责', 'qualification', '任职要求', '要求']:
                            if kw.lower() in next_text.lower():
                                still_title = True
                                break
                        if not still_title and len(next_text) < 10:
                            # 内容很少，可能是占位符，也当作内容框
                            content_row = r
                            break
                    
                    # 兜底：如果没找到空的，就用第2行下面的
                    if content_row is None and row_idx + 2 < rows_count:
                        content_row = row_idx + 2
                    
                    if content_row is not None:
                        target_cell = table.rows[content_row].cells[col_idx]
                        
                        if matched_content == 'responsibilities':
                            resp = data.get('responsibilities', [])
                            if isinstance(resp, str):
                                resp = [line.strip('•- ').strip() for line in resp.split('\n') if line.strip()]
                            _fill_cell_with_bullets(target_cell, resp)
                        
                        elif matched_content == 'qualifications':
                            items = []
                            if data.get('education'):
                                items.append(('教育背景：', data['education']))
                            if data.get('experience'):
                                items.append(('工作经验：', data['experience']))
                            if data.get('skills'):
                                items.append(('专业技能：', data['skills']))
                            if data.get('abilities'):
                                items.append(('能力要求：', data['abilities']))
                            _fill_cell_with_labeled_items(target_cell, items)
                        
                        else:
                            value = data.get(matched_content, '')
                            if value:
                                _set_cell_content(target_cell, str(value))

def _set_cell_content(cell, text):
    """设置单元格纯文本内容，清空原有内容，保留基础样式"""
    # 清空所有段落
    for p in cell.paragraphs:
        p.clear()
    
    # 使用第一个段落写入文本
    if cell.paragraphs:
        para = cell.paragraphs[0]
        para.text = text
    else:
        cell.text = text

def _fill_cell_with_bullets(cell, items):
    """用项目符号列表填充单元格，美化格式"""
    # 清空所有段落
    for p in cell.paragraphs:
        p.clear()
    
    if not items:
        return
    
    # 第一条用第一个段落
    first_para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    first_para.text = f"• {items[0]}"
    first_para.paragraph_format.left_indent = Cm(0.3)
    first_para.paragraph_format.space_after = Pt(2)
    
    # 其余条目新增段落
    for item in items[1:]:
        para = cell.add_paragraph()
        para.text = f"• {item}"
        para.paragraph_format.left_indent = Cm(0.3)
        para.paragraph_format.space_after = Pt(2)

def _fill_cell_with_labeled_items(cell, items):
    """用带标签的条目填充单元格（如 教育背景：xxx）"""
    # 清空所有段落
    for p in cell.paragraphs:
        p.clear()
    
    if not items:
        return
    
    # 第一条
    first_para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    label, value = items[0]
    run = first_para.add_run(label)
    run.bold = True
    first_para.add_run(value)
    first_para.paragraph_format.space_after = Pt(2)
    
    # 其余条目
    for label, value in items[1:]:
        para = cell.add_paragraph()
        run = para.add_run(label)
        run.bold = True
        para.add_run(value)
        para.paragraph_format.space_after = Pt(2)

def _fill_paragraph_sections(doc, data):
    """填充段落部分：岗位目的、主要职责、任职要求
    策略：找到各部分标题，定位内容区间，清空后插入新内容
    """
    
    paragraphs = doc.paragraphs
    
    # 第一步：找到所有章节标题的位置
    sections = {}  # name -> start_index
    
    for i, para in enumerate(paragraphs):
        text = para.text.strip().lower()
        
        # 岗位目的标题
        if any(kw in text for kw in ['position purpose', '岗位目的', '职位概述', 'position summary']):
            if 'purpose' not in sections:
                sections['purpose'] = i
        
        # 主要职责标题
        elif any(kw in text for kw in ['major task', '主要职责', '岗位职责', 'responsibilities', '工作描述']):
            if 'responsibilities' not in sections:
                sections['responsibilities'] = i
        
        # 任职要求标题
        elif any(kw in text for kw in ['qualification', '任职要求', '岗位要求', 'requirements']):
            if 'qualifications' not in sections:
                sections['qualifications'] = i
        
        # 审批栏/底部签名（作为结束标记）
        elif any(kw in text for kw in ['line manager', '签收人', 'approved by', '审批']):
            if 'end' not in sections:
                sections['end'] = i
    
    # 第二步：确定每个section的内容范围并填充
    
    # 1. 填充岗位目的
    if 'purpose' in sections:
        start = sections['purpose'] + 1
        end = sections.get('responsibilities', sections.get('end', start + 5))
        # 跳过副标题行（中文小标题）
        while start < end and paragraphs[start].text.strip():
            text = paragraphs[start].text.strip()
            if any(kw in text for kw in ['岗位目的', '表明该职位', '简要陈述', 'brief statement']):
                start += 1
            else:
                break
        # 跳空行
        while start < end and not paragraphs[start].text.strip():
            start += 1
        
        # 清空内容区域第一行，填入岗位目的
        purpose_text = data.get('position_purpose', '')
        if purpose_text and start < len(paragraphs):
            # 清空原有内容到下一个标题前
            _clear_paragraph_range(paragraphs, start, end)
            paragraphs[start].text = purpose_text
    
    # 2. 填充主要职责
    if 'responsibilities' in sections:
        start = sections['responsibilities'] + 1
        end = sections.get('qualifications', sections.get('end', start + 20))
        # 跳过副标题
        while start < end and paragraphs[start].text.strip():
            text = paragraphs[start].text.strip()
            if any(kw in text for kw in ['岗位主要任务', '主要任务与职责', 'position']):
                start += 1
            else:
                break
        # 跳空行
        while start < end and not paragraphs[start].text.strip():
            start += 1
        
        responsibilities = data.get('responsibilities', [])
        if isinstance(responsibilities, str):
            responsibilities = [line.strip('•- ').strip() for line in responsibilities.split('\n') if line.strip()]
        
        if responsibilities and start < len(paragraphs):
            # 清空原有内容区域
            _clear_paragraph_range(paragraphs, start, end)
            
            # 倒序插入（因为insert_paragraph_before是往前插）
            for resp in reversed(responsibilities):
                new_para = paragraphs[start].insert_paragraph_before(f"• {resp}")
                new_para.paragraph_format.left_indent = Cm(0.5)
    
    # 3. 填充任职要求
    if 'qualifications' in sections:
        start = sections['qualifications'] + 1
        end = sections.get('end', start + 20)
        # 跳过副标题
        while start < end and paragraphs[start].text.strip():
            text = paragraphs[start].text.strip()
            if any(kw in text for kw in ['教育，技能', '任职要求', '教育']):
                start += 1
            else:
                break
        # 跳空行
        while start < end and not paragraphs[start].text.strip():
            start += 1
        
        if start < len(paragraphs):
            # 清空原有内容区域
            _clear_paragraph_range(paragraphs, start, end)
            
            # 倒序插入各项
            items = []
            if data.get('education'):
                items.append(('教育背景：', data['education']))
            if data.get('experience'):
                items.append(('工作经验：', data['experience']))
            if data.get('skills'):
                items.append(('专业技能：', data['skills']))
            if data.get('abilities'):
                items.append(('能力要求：', data['abilities']))
            
            for label, value in reversed(items):
                p = paragraphs[start].insert_paragraph_before()
                run = p.add_run(label)
                run.bold = True
                p.add_run(value)

def _clear_paragraph_range(paragraphs, start, end):
    """清空指定范围内的段落文本"""
    for i in range(start, min(end, len(paragraphs))):
        paragraphs[i].text = ''
